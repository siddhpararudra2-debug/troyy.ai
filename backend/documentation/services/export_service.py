import json
import base64
import time
import markdown
from docx import Document
from documentation.services.template_engine import TemplateEngine
from sqlalchemy.orm import Session
from documentation.models.database import ProjectReport, CalculationReport

class ExportService:
    def __init__(self, db: Session):
        self.db = db
        self.template_engine = TemplateEngine()

    def export_report(self, report_id: int, export_format: str) -> dict:
        start_time = time.perf_counter()
        
        report = self.db.query(ProjectReport).filter(ProjectReport.id == report_id).first()
        if not report:
            raise ValueError("Report not found")
            
        calc = self.db.query(CalculationReport).filter(CalculationReport.project_report_id == report_id).first()
        
        # Prepare data dictionary
        data = {
            "title": report.title,
            "problem_statement": calc.problem_statement if calc else "",
            "requirements": calc.requirements if calc else "",
            "known_variables": calc.known_variables if calc else {},
            "unknown_variables": calc.unknown_variables if calc else {},
            "assumptions": calc.assumptions if calc else "",
            "formula_selection": calc.formula_selection if calc else "",
            "formula_explanation": calc.formula_explanation if calc else "",
            "unit_analysis": calc.unit_analysis if calc else "",
            "substitution_steps": calc.substitution_steps if calc else "",
            "intermediate_calculations": calc.intermediate_calculations if calc else {},
            "final_results": calc.final_results if calc else {},
            "verification_results": calc.verification_results if calc else "",
            "engineering_interpretation": calc.engineering_interpretation if calc else "",
            "recommendations": calc.recommendations if calc else ""
        }

        content = ""
        if export_format.upper() == "JSON":
            content = json.dumps(data, indent=2)
        elif export_format.upper() == "MARKDOWN":
            content = self.template_engine.render("AEROSPACE", "CALCULATION", data)
        elif export_format.upper() == "HTML":
            md_content = self.template_engine.render("AEROSPACE", "CALCULATION", data)
            content = markdown.markdown(md_content)
        elif export_format.upper() == "DOCX":
            content = self._generate_docx(data)
        else:
            raise ValueError(f"Unsupported format: {export_format}")

        execution_time = (time.perf_counter() - start_time) * 1000
        if execution_time > 2000:
            print(f"WARNING: Export took {execution_time}ms, exceeding 2s target")

        return {
            "report_id": report_id,
            "format": export_format.upper(),
            "content": content,
            "generated_at": report.generated_at
        }

    def _generate_docx(self, data: dict) -> str:
        """Lightweight DOCX generation to meet <2s requirement"""
        doc = Document()
        doc.add_heading(data.get("title", "Engineering Report"), 0)
        doc.add_heading("1. Problem Statement", level=1)
        doc.add_paragraph(data.get("problem_statement", ""))
        doc.add_heading("2. Formula & Analysis", level=1)
        doc.add_paragraph(f"Formula: {data.get('formula_selection', '')}")
        doc.add_paragraph(f"Unit Analysis: {data.get('unit_analysis', '')}")
        doc.add_heading("3. Results", level=1)
        doc.add_paragraph(str(data.get("final_results", {})))
        
        # Return as base64 string for API compatibility
        import io
        buffer = io.BytesIO()
        doc.save(buffer)
        return base64.b64encode(buffer.getvalue()).decode('utf-8')
